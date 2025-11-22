import os
import io
import re
import nltk
import spacy
from docx import Document
from pdfminer.high_level import extract_text
from nltk.corpus import stopwords
import traceback

# ----------------------------------------------------------
# 1. File Text Extraction (PDF / DOCX)
# ----------------------------------------------------------

def extract_text_from_file(filepath):
    """Extracts raw text content from PDF or DOCX file."""
    extension = os.path.splitext(filepath)[1].lower()

    if extension == ".pdf":
        try:
            # Try using extract_text with file path directly
            try:
                text = extract_text(filepath)
            except TypeError:
                # If that doesn't work, use file handle
                with open(filepath, "rb") as fh:
                    text = extract_text(fh)
            
            if text and text.strip():
                print(f"PDF extraction successful. Extracted {len(text)} characters.")
                return text
            else:
                print("PDF extraction returned empty text")
                return None
        except Exception as e:
            print(f"PDF extraction error: {e}")
            print(traceback.format_exc())
            return None

    elif extension == ".docx":
        try:
            doc = Document(filepath)
            # Extract text from paragraphs
            text = [paragraph.text for paragraph in doc.paragraphs]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            result = "\n".join(text)
            return result if result.strip() else None
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            print(traceback.format_exc())
            return None

    return None


# ----------------------------------------------------------
# 2. NLP RESOURCE SETUP
# ----------------------------------------------------------

# Fix: Correct NLTK resource handling
try:
    nltk.data.find("corpora/stopwords")
except LookupError:
    nltk.download("stopwords")

# Fix: spaCy safe load
nlp = None
try:
    nlp = spacy.load("en_core_web_sm")
except (IOError, OSError):
    # If missing, download automatically
    try:
        import subprocess
        import sys
        subprocess.run([sys.executable, "-m", "spacy", "download", "en_core_web_sm"], check=True)
        nlp = spacy.load("en_core_web_sm")
    except Exception as e:
        print(f"Warning: Could not load spaCy model: {e}")
        print("Continuing without spaCy NLP features...")
        nlp = None

STOPWORDS = set(stopwords.words("english"))

TARGET_JOB_KEYWORDS = {
    "developer": ["Python", "Flask", "React", "TypeScript", "SQL", "Git"],
    "data_science": ["Python", "Pandas", "NumPy", "TensorFlow", "PyTorch", "Stats"]
}

JOB_ROLE = "developer"  # Hardcoded for now


# ----------------------------------------------------------
# 3. Resume Analyzer
# ----------------------------------------------------------
# ANALYSIS CRITERIA:
# 1. KEYWORD OPTIMIZATION (40 points): Checks for industry-relevant keywords
#    - Each matched keyword adds 8 points (max 40 for 5 keywords)
#    - Missing keywords are suggested for improvement
#
# 2. STRUCTURAL COMPLETENESS (30 points): Ensures essential sections exist
#    - Skills section (10 points)
#    - Experience/Work History section (10 points)
#    - Education section (5 points)
#    - Contact information (5 points)
#
# 3. CONTENT QUALITY (20 points): Evaluates resume depth and clarity
#    - Appropriate length (not too short, not too long)
#    - Action verbs usage
#    - Quantifiable achievements
#
# 4. ATS COMPATIBILITY (10 points): Format and parsing considerations
#    - Text-based format (not image-based)
#    - Standard section headers
#    - Proper formatting

def analyze_resume_text(raw_text, job_keywords):
    """
    Analyzes resume text based on ATS (Applicant Tracking System) criteria.
    
    Analysis Basis:
    - Keyword Optimization: Matches resume against target job keywords
    - Structural Completeness: Checks for essential resume sections
    - Content Quality: Evaluates depth, clarity, and impact
    - ATS Compatibility: Ensures resume can be parsed by ATS systems
    
    Returns:
        dict with atsScore, keywordSuggestions, and formattingTips
    """
    if not raw_text:
        return {
            "atsScore": 0,
            "keywordSuggestions": job_keywords.copy() if job_keywords else [],
            "formattingTips": ["Could not read file text. Please ensure the file is not corrupted and is a valid PDF or DOCX format."]
        }

    text_lower = raw_text.lower()
    text_words = raw_text.split()
    word_count = len(text_words)
    
    # Use spaCy if available, otherwise skip NLP features
    doc = None
    if nlp:
        try:
            doc = nlp(raw_text)
        except Exception as e:
            print(f"spaCy processing error: {e}")
            doc = None

    score = 0
    matched_keywords = []
    tips = []

    # ============================================
    # 1. KEYWORD OPTIMIZATION (40 points max)
    # ============================================
    keyword_score = 0
    max_keyword_points = min(40, len(job_keywords) * 8)
    points_per_keyword = max_keyword_points / len(job_keywords) if job_keywords else 0
    
    for kw in job_keywords:
        # Check for exact keyword match (case-insensitive)
        if kw.lower() in text_lower:
            keyword_score += points_per_keyword
            matched_keywords.append(kw)
        # Also check for partial matches (e.g., "Python" matches "Python3", "Pythonic")
        elif any(kw.lower() in word.lower() for word in text_words if len(word) > 3):
            keyword_score += points_per_keyword * 0.5  # Partial credit
            matched_keywords.append(kw)
    
    score += int(keyword_score)
    
    # ============================================
    # 2. STRUCTURAL COMPLETENESS (30 points max)
    # ============================================
    structural_score = 0
    
    # Skills section check (10 points)
    skills_patterns = [r"\bskills\b", r"\btechnical\s+skills\b", r"\bcompetencies\b", r"\bproficiencies\b"]
    has_skills = any(re.search(pattern, text_lower) for pattern in skills_patterns)
    if has_skills:
        structural_score += 10
    else:
        tips.append("Add a dedicated 'Skills' or 'Technical Skills' section to highlight your competencies.")
    
    # Experience section check (10 points)
    experience_patterns = [r"\bexperience\b", r"\bwork\s+history\b", r"\bemployment\b", r"\bprofessional\s+experience\b", r"\bwork\s+experience\b"]
    has_experience = any(re.search(pattern, text_lower) for pattern in experience_patterns)
    if has_experience:
        structural_score += 10
    else:
        tips.append("Add an 'Experience', 'Work History', or 'Professional Experience' section.")
    
    # Education section check (5 points)
    education_patterns = [r"\beducation\b", r"\bacademic\b", r"\bqualifications\b"]
    has_education = any(re.search(pattern, text_lower) for pattern in education_patterns)
    if has_education:
        structural_score += 5
    else:
        tips.append("Include an 'Education' section with your academic qualifications.")
    
    # Contact information check (5 points)
    contact_patterns = [r"@", r"\bemail\b", r"\bphone\b", r"\bmobile\b", r"\bcontact\b"]
    has_contact = any(re.search(pattern, text_lower) for pattern in contact_patterns)
    if has_contact:
        structural_score += 5
    else:
        tips.append("Ensure your contact information (email, phone) is clearly visible.")
    
    score += structural_score

    # ============================================
    # 3. CONTENT QUALITY (20 points max)
    # ============================================
    content_score = 0
    
    # Length check (10 points)
    if 200 <= word_count <= 800:  # Ideal range: 200-800 words
        content_score += 10
    elif word_count < 200:
        content_score += 5
        tips.append(f"Your resume is quite short ({word_count} words). Add more detail about your experience, projects, and achievements.")
    elif word_count > 800:
        content_score += 5
        tips.append(f"Your resume is lengthy ({word_count} words). Consider condensing to 1-2 pages for better readability.")
    else:
        content_score += 3
    
    # Action verbs check (5 points) - Strong action verbs indicate impact
    action_verbs = ["developed", "created", "implemented", "designed", "managed", "led", "improved", 
                    "achieved", "optimized", "built", "delivered", "executed", "launched", "established"]
    action_verb_count = sum(1 for verb in action_verbs if verb in text_lower)
    if action_verb_count >= 5:
        content_score += 5
    elif action_verb_count >= 3:
        content_score += 3
        tips.append("Use more action verbs (e.g., 'developed', 'created', 'implemented') to make your achievements stand out.")
    else:
        tips.append("Include more action verbs to describe your accomplishments and responsibilities.")
    
    # Quantifiable achievements check (5 points)
    number_patterns = [r"\d+%", r"\d+\+", r"\$\d+", r"\d+\s+(years?|months?)", r"\d+\s+(people|users|customers)"]
    has_numbers = any(re.search(pattern, text_lower) for pattern in number_patterns)
    if has_numbers:
        content_score += 5
    else:
        tips.append("Add quantifiable metrics (percentages, numbers, timeframes) to demonstrate your impact.")
    
    score += content_score

    # ============================================
    # 4. ATS COMPATIBILITY (10 points max)
    # ============================================
    ats_score = 10  # Base score - if text was extracted, it's likely ATS-compatible
    ats_tips = []
    
    # Check for common ATS-unfriendly elements
    if re.search(r"\.(jpg|jpeg|png|gif)", text_lower):
        ats_tips.append("Avoid embedding images in your resume. ATS systems cannot read text from images.")
    
    # Check for proper section headers
    section_headers = ["summary", "objective", "skills", "experience", "education", "projects"]
    found_headers = sum(1 for header in section_headers if re.search(rf"\b{header}\b", text_lower))
    if found_headers < 3:
        ats_tips.append("Use clear, standard section headers (e.g., 'Experience', 'Education', 'Skills') for better ATS parsing.")
    
    if ats_tips:
        tips.extend(ats_tips)
    
    score += ats_score

    # ============================================
    # FINAL SCORE CALCULATION
    # ============================================
    # Normalize to 0-100 scale
    max_possible_score = 100  # 40 + 30 + 20 + 10
    ats_final_score = min(100, max(0, int((score / max_possible_score) * 100)))
    
    # Ensure we always return lists
    keyword_suggestions = [kw for kw in job_keywords if kw not in matched_keywords]
    
    # Add positive feedback if score is high
    if ats_final_score >= 80:
        if not any("good" in tip.lower() or "great" in tip.lower() for tip in tips):
            tips.insert(0, "Excellent resume! Your resume shows strong ATS compatibility and structure.")
    elif ats_final_score >= 60:
        tips.insert(0, "Your resume has a solid foundation. Consider the suggestions below to improve your ATS score further.")
    
    # If no tips, add a generic positive message
    if not tips:
        tips.append("Your resume structure looks good! Keep up the great work.")

    return {
        "atsScore": ats_final_score,
        "keywordSuggestions": keyword_suggestions,
        "formattingTips": tips
    }


# ----------------------------------------------------------
# 4. Entry Function
# ----------------------------------------------------------

def process_resume_file(filepath):
    """Main entry for resume analysis."""
    try:
        raw_text = extract_text_from_file(filepath)
        if raw_text is None:
            return {
                "atsScore": 0,
                "keywordSuggestions": TARGET_JOB_KEYWORDS.get(JOB_ROLE, TARGET_JOB_KEYWORDS["developer"]).copy(),
                "formattingTips": ["Failed to extract text from the file. Please ensure the file is not corrupted and is a valid PDF or DOCX format."]
            }
        
        keywords = TARGET_JOB_KEYWORDS.get(JOB_ROLE, TARGET_JOB_KEYWORDS["developer"])
        result = analyze_resume_text(raw_text, keywords)
        return result
    except Exception as e:
        print(f"Error in process_resume_file: {e}")
        print(traceback.format_exc())
        return {
            "atsScore": 0,
            "keywordSuggestions": TARGET_JOB_KEYWORDS.get(JOB_ROLE, TARGET_JOB_KEYWORDS["developer"]).copy(),
            "formattingTips": [f"An error occurred during analysis: {str(e)}"]
        }
